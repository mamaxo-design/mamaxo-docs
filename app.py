#!/usr/bin/env python3
"""
mamaXO Document Factory — web app (flat, self-contained for easy hosting).
Colleagues sign in with a shared password, fill the form, get a branded PDF.
Design is locked (base.css + brand.config.json); users only supply content.
All images/fonts are embedded (assets_b64.json / base.css) so there are no
sub-folders — the whole app is a flat set of files.
"""
import os, io, json, html, base64
from flask import Flask, request, session, redirect, send_file, render_template, Response
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
CFG = json.load(open(os.path.join(BASE, "brand.config.json"), encoding="utf-8"))
PROPS = json.load(open(os.path.join(BASE, "properties.json"), encoding="utf-8"))
ASSETS = json.load(open(os.path.join(BASE, "assets_b64.json"), encoding="utf-8"))

APP_PASSWORD = os.environ.get("APP_PASSWORD", "mamaxo2026")
SECRET = os.environ.get("SECRET_KEY", "change-me-in-production")

app = Flask(__name__, template_folder=".")
app.secret_key = SECRET

def esc(s): return html.escape(str(s if s is not None else ""))
def A(k): return ASSETS.get(k, "")

ICONS = {
 "eye": '<path d="M2 12s3.5-7 10-7 10 7 10 7-3.5 7-10 7S2 12 2 12z"/><circle cx="12" cy="12" r="3"/>',
 "people": '<circle cx="9" cy="8" r="3"/><path d="M3.5 20c0-3 2.7-5 5.5-5s5.5 2 5.5 5"/><path d="M16 6a3 3 0 0 1 0 6"/><path d="M16.5 15c2.4.2 4.5 2 4.5 5"/>',
 "globe": '<circle cx="12" cy="12" r="9"/><path d="M3 12h18"/><path d="M12 3c3 3 3 15 0 18"/><path d="M12 3c-3 3-3 15 0 18"/>',
 "building": '<rect x="4" y="3" width="10" height="18"/><path d="M14 8h6v13H4"/><path d="M7 7h2M7 11h2M7 15h2"/>',
 "laptop": '<rect x="3" y="4" width="18" height="12" rx="1.5"/><path d="M8 20h8M12 16v4"/>',
 "house": '<path d="M4 11l8-6 8 6"/><path d="M6 10v10h12V10"/><rect x="10" y="14" width="4" height="6"/>',
 "flow": '<rect x="3" y="4" width="7" height="7" rx="1"/><rect x="14" y="13" width="7" height="7" rx="1"/><path d="M10 7h4a3 3 0 0 1 3 3v3"/>',
 "check": '<path d="M4 12l5 5L20 6"/>', "cross": '<path d="M6 6l12 12M18 6L6 18"/>',
}
DIFF_ICONS = ["building", "laptop", "house", "flow"]
def svg(icon, stroke="#E2681A", w="1.7"):
    return f'<svg viewBox="0 0 24 24" fill="none" stroke="{stroke}" stroke-width="{w}" stroke-linecap="round" stroke-linejoin="round">{ICONS[icon]}</svg>'

def cover(profile, data, lab):
    m = lambda k, v: f'<div><div class="k">{esc(lab.get(k,k))}</div><div class="v">{esc(v) or "—"}</div></div>'
    return f'''<section class="cover"><div class="gradbar full"></div>
      <div class="brandbar"><img src="{A('logo_full')}" alt="mamaXO"></div>
      <img class="badge" src="{A('logo_mark')}" alt="">
      <div class="cover-body"><div class="kicker">{esc(profile["kicker"])}</div>
        <h1>{esc(data.get("title") or profile["line_title"])}</h1><div class="cover-rule"></div>
        <p class="sub">{esc(data.get("subtitle",""))}</p>
        <div class="meta">{m("prepared_for",data.get("client"))}{m("date",data.get("date"))}
          <div><div class="k">{esc(lab.get("prepared_by","Prepared by"))}</div><div class="v">mamaXO</div></div>
          {m("ref",data.get("ref"))}</div></div></section>'''

def catalogue_page(data, lab):
    sel = data.get("selected") or []
    chosen = [PROPS[i] for i in sel if 0 <= i < len(PROPS)] or PROPS[:4]
    frm = esc(lab.get("from", "from"))
    cards = ""
    for p in chosen:
        key = os.path.splitext(p["img"])[0]
        tag = "tag" if p["status"] == "Available" else "tag soon"
        cards += f'''<article class="pcard"><div class="ph" style="background-image:url('{A(key)}')">
          <span class="{tag}">{esc(p["status"])}</span></div><div class="bd"><div class="loc">{esc(p["loc"])}</div>
          <h3>{esc(p["name"])}</h3><div class="price"><span class="from">{frm}</span>€{p["price"]:,}</div>
          <div class="specs"><span><b>{esc(p["beds"])}</b> bd</span><span><b>{esc(p["baths"])}</b> ba</span><span><b>{esc(p["area"])}</b> m²</span></div>
          <p>{esc(p["desc"])}</p></div></article>'''
    return f'''<div class="page"><div class="snum"><span class="n">01</span><span class="kick">Selected properties</span></div>
      <h2>{esc(data.get("title") or "Selected properties")}</h2><hr class="rule">
      <p class="lead">{esc(data.get("lead","A shortlist matched to your brief. Each property is verified before anyone commits."))}</p>
      <div class="grid">{cards}</div></div>'''

def content_page(data):
    parts = []
    lead = (data.get("lead") or "").strip()
    if lead:
        parts.append(f'<p class="lead">{esc(lead)}</p>')
    kpis = data.get("kpis") or []
    if kpis:
        kh = "".join(f'<div class="kpi"><div class="n">{esc(k[0])}</div><div class="d">{esc(k[1])}</div></div>' for k in kpis if k and k[0])
        if kh:
            parts.append(f'<div class="kpis">{kh}</div>')
    for s in (data.get("sections") or []):
        title = (s.get("title") or "").strip()
        body = (s.get("text") or "").strip()
        if not title and not body:
            continue
        if title:
            parts.append(f'<h2>{esc(title)}</h2><hr class="rule">')
        for para in body.split("\n"):
            if para.strip():
                parts.append(f'<p>{esc(para.strip())}</p>')
    if not parts:
        parts.append('<p class="lead" style="opacity:.5">No content yet — add sections in the form.</p>')
    return f'''<div class="page"><div class="snum"><span class="n">01</span><span class="kick">{esc(data.get("title") or "Document")}</span></div>
      {''.join(parts)}</div>'''

def team_html(profile):
    cells = ""
    for m in profile.get("team", []):
        av = (f'<img class="av" src="{A(m["photo"])}">' if m.get("photo") and A(m["photo"])
              else f'<div class="av">{esc(m.get("initials","•"))}</div>')
        cells += f'<div class="person">{av}<div class="nm">{esc(m["name"])}</div><div class="rl">{esc(m["role"])}</div></div>'
    return cells

def about_page(profile, line):
    if line == "golden_visa":
        eyebrow, title = "About us", "Why mamaXO"
        sub = "An Athens-based team for international investors — local expertise, remote simplicity, support from purchase to approval."
        cards = "".join(f'<div class="card"><div class="ico">{svg(DIFF_ICONS[i%4])}</div><h3>{esc(d["title"])}</h3><p>{esc(d["text"])}</p></div>'
                        for i, d in enumerate(CFG.get("difference", [])))
        block = f'<div class="label">The mamaXO difference</div><div class="feat">{cards}</div>'
    else:
        eyebrow, title, sub = "What we stand for", "Our values", ""
        cards = "".join(f'<div class="card"><div class="ico">{svg(v.get("icon","globe"))}</div><h3>{esc(v["title"])}</h3><p>{esc(v["text"])}</p></div>'
                        for v in profile.get("values", []))
        block = f'<div class="feat three">{cards}</div>'
    return f'''<div class="page"><div class="about"><div class="eyebrow">{esc(eyebrow)}</div>
      <h2>{esc(title)}</h2><hr class="rule">{f'<p class="sub">{esc(sub)}</p>' if sub else ''}{block}
      <div class="label">The team</div><div class="team">{team_html(profile)}</div></div></div>'''

def secondary_page(profile, line):
    if line == "golden_visa":
        cards = "".join(f'<div class="card"><div class="n">{i:02d}</div><h3>{esc(s["title"])}</h3><p>{esc(s["text"])}</p></div>'
                        for i, s in enumerate(CFG.get("services_after_purchase", []), 1))
        return f'''<section class="services"><div class="eyebrow">After the purchase</div>
          <h2>We stay with you</h2><hr class="rule"><p class="sub">Full support once you own — so ownership stays effortless.</p>
          <div class="svc">{cards}</div></section>'''
    cmp = profile.get("comparison", {})
    us = "".join(f'<li>{svg("check","#E2681A","2.2")}{esc(x)}</li>' for x in cmp.get("us", []))
    them = "".join(f'<li>{svg("cross","#B7B0A4","2.2")}{esc(x)}</li>' for x in cmp.get("them", []))
    return f'''<div class="page"><h2 style="text-align:center;font-size:22pt;margin-bottom:18px">{esc(cmp.get("title","Better than a private landlord."))}</h2>
      <div class="compare"><div class="col us"><h3>mamaXO</h3><ul>{us}</ul></div>
      <div class="col them"><h3>{esc(cmp.get("them_label","Alternative"))}</h3><ul>{them}</ul></div></div></div>'''

def plate_page(profile, lab, disclaimer, variant):
    photo = f'<div class="bg" style="background-image:url(\'{A("plate")}\')"></div>' if variant == "photo" else ""
    cls = "plate photo" if variant == "photo" else "plate"
    return f'''<section class="{cls}"><div class="card">{photo}
      <div class="eyebrow">{esc(lab.get("next_step","Next step"))}</div><h2>{esc(profile["line_title"])}</h2>
      <p>{esc(profile["plate_blurb"])}</p><p>{esc(profile["plate_line2"])}</p>
      <div class="pill">{esc(profile["email"])} &nbsp;·&nbsp; {esc(profile["phone"])}</div>
      <div class="addr">{esc(profile["address"])}</div>
      <div class="addr" style="opacity:.7;margin-top:10px;font-size:7.5pt">{esc(disclaimer)}</div></div></section>'''

def build_document(data):
    line = data.get("line", "golden_visa")
    profile = CFG["profiles"][line]
    lang = data.get("lang", "en")
    lab = CFG.get("languages", {}).get("labels", {}).get(lang, {})
    disclaimer = CFG.get("disclaimers", {}).get(lang, CFG.get("disclaimers", {}).get("en", ""))
    ends = data.get("ends", {})
    parts = [cover(profile, data, lab)]
    parts.append(catalogue_page(data, lab) if data.get("type") == "catalogue" else content_page(data))
    if ends.get("about"): parts.append(about_page(profile, line))
    if ends.get("svc"):   parts.append(secondary_page(profile, line))
    parts.append(plate_page(profile, lab, disclaimer, ends.get("plate", "orange")))
    return f'''<!DOCTYPE html><html lang="{esc(lang)}"><head><meta charset="utf-8">
      <link rel="stylesheet" href="base.css"><title>{esc(data.get("title") or "mamaXO")}</title></head><body>
      <div class="doc-meta-footer">{esc(data.get("title") or profile["line_title"])}</div>{''.join(parts)}</body></html>'''

# ----------------------------------------------------------------------------- Word (.docx) export
_ORANGE = RGBColor(0xE2, 0x68, 0x1A)
_INK = RGBColor(0x2B, 0x2B, 0x2D)
_MUT = RGBColor(0x7A, 0x75, 0x6B)

def _img_stream(datauri):
    return io.BytesIO(base64.b64decode(datauri.split(",", 1)[1]))

def _dx_heading(doc, text, size=18):
    h = doc.add_paragraph(); r = h.add_run(text)
    r.bold = True; r.font.name = "Georgia"; r.font.size = Pt(size); r.font.color.rgb = _INK

def build_docx(data):
    line = data.get("line", "golden_visa")
    profile = CFG["profiles"][line]
    lang = data.get("lang", "en")
    lab = CFG.get("languages", {}).get("labels", {}).get(lang, {})
    disclaimer = CFG.get("disclaimers", {}).get(lang, CFG.get("disclaimers", {}).get("en", ""))

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = _INK

    p = doc.add_paragraph(); r = p.add_run(profile["kicker"].upper())
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _ORANGE
    h = doc.add_paragraph(); rr = h.add_run(data.get("title") or profile["line_title"])
    rr.bold = True; rr.font.name = "Georgia"; rr.font.size = Pt(26); rr.font.color.rgb = _INK

    meta = []
    if data.get("client"): meta.append(f'{lab.get("prepared_for","Prepared for")}: {data["client"]}')
    if data.get("date"):   meta.append(f'{lab.get("date","Date")}: {data["date"]}')
    if data.get("ref"):    meta.append(f'{lab.get("ref","Ref")}: {data["ref"]}')
    if meta:
        mp = doc.add_paragraph(); mr = mp.add_run("   ·   ".join(meta)); mr.font.size = Pt(9); mr.font.color.rgb = _MUT
    doc.add_paragraph()

    if (data.get("lead") or "").strip():
        lp = doc.add_paragraph(); lr = lp.add_run(data["lead"].strip()); lr.font.size = Pt(12); lr.italic = True
    for s in (data.get("sections") or []):
        title = (s.get("title") or "").strip(); body = (s.get("text") or "").strip()
        if not title and not body: continue
        if title: _dx_heading(doc, title, size=15)
        for para in body.split("\n"):
            if para.strip(): doc.add_paragraph(para.strip())

    doc.add_page_break()
    if line == "golden_visa":
        _dx_heading(doc, "Why mamaXO")
        for d in CFG.get("difference", []):
            b = doc.add_paragraph(style="List Bullet"); br = b.add_run(d["title"] + " — "); br.bold = True; b.add_run(d["text"])
    else:
        _dx_heading(doc, "Our values")
        for v in profile.get("values", []):
            b = doc.add_paragraph(style="List Bullet"); br = b.add_run(v["title"] + " — "); br.bold = True; b.add_run(v["text"])

    tp = doc.add_paragraph(); tr = tp.add_run("THE TEAM"); tr.bold = True; tr.font.size = Pt(9); tr.font.color.rgb = _ORANGE
    team = profile.get("team", [])
    cols = 4
    table = doc.add_table(rows=0, cols=cols)
    for r0 in range((len(team) + cols - 1) // cols):
        cells = table.add_row().cells
        for c0 in range(cols):
            idx = r0 * cols + c0
            if idx >= len(team): continue
            m = team[idx]; cell = cells[c0]
            cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            key = m.get("photo")
            if key and ASSETS.get(key):
                try: cp.add_run().add_picture(_img_stream(ASSETS[key]), width=Inches(1.05))
                except Exception: pass
            npa = cell.add_paragraph(); npa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nr = npa.add_run(m["name"]); nr.bold = True; nr.font.size = Pt(9)
            rpa = cell.add_paragraph(); rpa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = rpa.add_run(m["role"]); r2.font.size = Pt(8); r2.font.color.rgb = _ORANGE

    doc.add_paragraph()
    _dx_heading(doc, profile["line_title"], size=15)
    cp = doc.add_paragraph(); cp.add_run(f'{profile["email"]}   ·   {profile["phone"]}\n{profile["address"]}')
    dp = doc.add_paragraph(); dr = dp.add_run(disclaimer); dr.font.size = Pt(8); dr.font.color.rgb = _MUT

    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        if request.form.get("password") == APP_PASSWORD:
            session["ok"] = True
            return redirect("/")
        return render_template("login.html", error=True)
    return render_template("login.html", error=False)

@app.route("/logout")
def logout():
    session.clear(); return redirect("/login")

@app.route("/")
def index():
    if not session.get("ok"): return redirect("/login")
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    if not session.get("ok"): return Response("unauthorized", 401)
    data = request.get_json(force=True)
    base_name = (data.get("client") or "document").strip().replace(" ", "_") + "_" + data.get("type", "doc")
    if data.get("format") == "docx":
        bio = build_docx(data)
        return send_file(bio, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name=base_name + ".docx")
    pdf = HTML(string=build_document(data), base_url=BASE).write_pdf()
    return send_file(io.BytesIO(pdf), mimetype="application/pdf", as_attachment=True, download_name=base_name + ".pdf")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 7860)), debug=False)
